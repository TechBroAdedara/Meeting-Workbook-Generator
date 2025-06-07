import asyncio
import re
import time
from collections.abc import Sequence
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Inches
from docx.table import Table, _Cell

from MyScrapper import parse_workbook_url

# record = {'date': 'OCTOBER 27–NOVEMBER\xa02',
#           'main_table': ['TREASURES FROM GOD’S WORD',
#                          (1, 'Enjoy a Happy, Healthy Life (10 min.)'),
#                          (2, 'Spiritual Gems (10 min.)'),
#                          (3, 'Bible Reading (4 min.)'),
#                          'APPLY YOURSELF TO THE FIELD MINISTRY',
#                          (4, 'Following Up (3 min.)'),
#                          (5, 'Following Up (4 min.)'),
#                          (6, 'Talk (5 min.)'),
#                          'LIVING AS CHRISTIANS',
#                          (7, 'Local Needs (15 min.)'),
#                          (8, 'Congregation Bible Study (30 min.)')]}

document = Document()

sections = document.sections
for section in sections:
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def rgb_to_hex(rgb_color: tuple):
    """Accepts (R, G, B) and returns hex string like 'FF0000'."""
    return '{:02X}{:02X}{:02X}'.format(*rgb_color)


def set_cell_background(cell, rgb_color):
    """
    Low level cell background color change.
    Python-docx does not provide a high-level way to edit cell background.
    """
    hex_color = rgb_to_hex(rgb_color)

    # Get or add table cell properties
    cell_properties = cell._element.get_or_add_tcPr()

    # Create shading element
    cell_shading = OxmlElement("w:shd")
    cell_shading.set(qn("w:val"), "clear")  # or "solid"
    cell_shading.set(qn("w:color"), "auto")  # foreground color
    cell_shading.set(qn("w:fill"), hex_color)  # background color

    # Append shading to cell properties
    cell_properties.append(cell_shading)


def add_cell_text(cell: _Cell,
                  text: str,
                  text_color: RGBColor = RGBColor(255, 255, 255),
                  font_name: str = "Calibri",
                  font_size: int = 12,
                  bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run(text)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.color.rgb = text_color


def add_custom_row(word_table: Table, no_of_cols: int, divisions: list[float], ) -> Sequence[_Cell]:
    """
    Adds row with custom partition to table.
    'Custom' in the sense that it creates a row with specified partitions (based on percentages) and returns each of the partitions.
    (Very proud of this function btw.)
    
    :param word_table: The table to create row in.
    :param no_of_cols: Number of columns in the table.
    :param divisions: List of division percentages(in respective order). Must sum up to 1.

    :returns: sequence - Cell partitions of the row
    """
    if sum(divisions) > 1:
        raise ValueError(f"The sum of divisions should not be greater than 1 (100%). Current sum: {sum(divisions)} ")

    if any([div < 0 for div in divisions]):
        raise ValueError("You cannot have a negative division")

    partitions_idx = []
    cell_begin = 0
    table_row_cells = word_table.add_row().cells

    count = 0
    for div in divisions:
        # If the division we are on is not the last one, set ``cell end`` to be the computed end of merged cells.
        # else make the ``cell_end`` the last cell in the row
        cell_end = round(cell_begin + (div * no_of_cols) - 1) \
            if count != len(divisions) - 1 \
            else (no_of_cols - 1)

        table_row_cells[cell_begin].merge(table_row_cells[cell_end])
        partitions_idx.append(cell_begin)
        cell_begin = cell_end + 1  # The next cell to be merged to be after the previous merged cell's end.
        count += 1

    return [table_row_cells[partition] for partition in partitions_idx]


async def generate_word_document():
    cols = 24

    """Regex for formatting third cell conditionally"""
    partner_assignment_pattern = re.compile("Starting a Con|Following Up|Explaining Your|Making Disciples")
    single_assignment_pattern = re.compile(r"Bible Reading|\w+—Imitate \w+|\w+—What \w+ Did|Talk")

    # Performance counter
    start = time.perf_counter()

    # User input for grabbing contents of workbook
    month: Literal["January", "March", "May", "July", "September", "November"] = input(  # type: ignore
        "Enter the month you want to get workbook contents for? "
        " \n Valid months: [\"January\", \"March\", \"May\", \"July\", \"September\", \"November\"] \n")
    list_of_possible_months = ["january", "march", "may", "july", "september", "november"]

    if month.lower() not in list_of_possible_months:
        raise ValueError("Invalid month. Please input a month from the valid list of months.")
    # Get the cover title and workbook content from my web scrapper
    cover_title, workbook_contents = await parse_workbook_url(month)

    # For aesthetic purposes. Useful for tracking progress of word generator
    total_progress = len(workbook_contents)
    count = 0

    print("\n\n")
    print("<-----------------WORD DOCUMENT PROCESSING----------------------->")

    heading = document.add_heading("AKEJA CONGREGATION MIDWEEK MEETING SCHEDULE", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align to the center
    heading.paragraph_format.space_after = Pt(4)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black Text

    subheading = document.add_heading(cover_title, level=2)  # The dynamic text for each month
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align to the center
    subheading.paragraph_format.space_before = Pt(2)
    subheading.paragraph_format.space_after = Pt(10)

    for record in workbook_contents:
        if record is None:
            print(f"SKIPPED : {record}")
            continue # If there is an error within a week's content, skip it

        new_table = document.add_table(cols=cols, rows=0)  # Add a new table for each run

        """ Partition the header row containing the date and get partitions back """
        hdr_partitions = add_custom_row(new_table, cols, [1])  # One partition
        add_cell_text(cell=hdr_partitions[0],
                      text=str(record["date"]),
                      bold=True,
                      text_color=RGBColor(0, 0, 0))

        """The first row of the table, consistent across tables. """
        first_row = add_custom_row(new_table, cols, divisions=[0.8, 0.2])  # Two partitions
        add_cell_text(cell=first_row[0],
                      text="Chairman's Opening Comments - 1 min & Prayer",
                      text_color=RGBColor(0, 0, 0),
                      bold=True)
        add_cell_text(cell=first_row[1],
                      text="Stephen Adeloro",
                      font_size=11,
                      text_color=RGBColor(0, 0, 0))

        # Populating table with the workbook content.
        is_after_living_as_christians = False

        # text_is_short = lambda r: len(str(r[1])) < 35
        for row in record["main_table"]:
            if row == "LIVING AS CHRISTIANS":
                is_after_living_as_christians = True

            if type(row) is not tuple:
                """ If the row is a meeting part e.g. Treasures From God's Word."""
                rgb_color = (98, 101, 104) if str(row) == "TREASURES FROM GOD’S WORD" \
                    else (189, 142, 22) if str(row) == "APPLY YOURSELF TO THE FIELD MINISTRY" \
                    else (148, 54, 52)  # Conditional cell background based on different meeting parts.

                # Partition each row according to my percentages and get each partition back
                subheading_row = add_custom_row(new_table, cols, [0.7, 0.3])  # Two partitions

                set_cell_background(subheading_row[0], rgb_color=rgb_color)
                add_cell_text(subheading_row[0], str(row), bold=True, font_size=11)  # Adding text with custom styling
            else:
                """ If the row is a meeting subpart"""
                # Partition each row according to my percentages and get each partition back
                meeting_part_row = add_custom_row(new_table,
                                                  no_of_cols=cols,
                                                  divisions=(
                                                      [0.1, 0.45, 0.15, 0.3]
                                                      if not (is_after_living_as_christians or str(row[0]) in {"1",
                                                                                                               "2"})
                                                      else [0.1, 0.6, 0.3]
                                                  ))

                """First col in row"""
                add_cell_text(
                    meeting_part_row[0],
                    str(row[0]),
                    RGBColor(0, 0, 0),
                    font_size=9
                )  # Adding text with custom color

                """Second col in row"""
                add_cell_text(
                    meeting_part_row[1],
                    str(row[1]),
                    RGBColor(0, 0, 0),
                    font_size=11 if len(str(row[1])) < 50 else 10  # decrease font size if text is too long
                )  # Adding text with custom color

                """Third col in row. Dynamically render the third col according to the type of assignment"""
                student_partner_text = "Student/Partner" \
                    if partner_assignment_pattern.match(str(row[1])) \
                    else "Student" if single_assignment_pattern.match(str(row[1])) \
                    else ""

                add_cell_text(
                    meeting_part_row[2],
                    text=student_partner_text,
                    text_color=RGBColor(0, 0, 0),
                    font_size=9,
                )

        """The last row of the table, consistent across all tables"""
        last_row = add_custom_row(new_table, cols, [0.1, 0.6, 0.3])
        add_cell_text(last_row[1],
                      text="• Review/Preview/Announcements - (3 min.) ",
                      bold=True,
                      text_color=RGBColor(0, 0, 0),
                      font_size=11)

        count += 1  # Update progress tracker

        if count % 2 == 1:  # Add a page break to every odd-numbered table
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)

        if count % 2 == 0:  # Add a page break after every even-numbered table
            document.add_page_break()

        print(f"\rProcessed word document for {record["date"]} - Progress: {round((count / total_progress) * 100)}%",
              end=" ", flush=True)

    print("\n\nWord document has finished processing.")
    end = time.perf_counter()
    print(f"Ran for {end - start:.4f} seconds")

    document.save(f"GeneratedDocs/{cover_title}.docx")


if __name__ == "__main__":
    asyncio.run(generate_word_document())
