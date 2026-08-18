import asyncio
import io
import re
import time
from typing import Annotated, Literal

import aiohttp
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from fastapi import FastAPI, Form, HTTPException, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from collections.abc import Sequence

from icecream import ic

# ---------------------------------------------------------------------------
# APP
# ---------------------------------------------------------------------------

app = FastAPI(title="Akeja Meeting Schedule Generator")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# LANGUAGE CONFIG
# ---------------------------------------------------------------------------

LANGUAGE_CONFIG = {
    "english": {
        "base_url": "https://wol.jw.org/",
        "workbook_path": "en/wol/library/r1/lp-e/all-publications/meeting-workbooks/life-and-ministry-meeting-workbook-2026/{month}",
        "song_pattern": re.compile(r"\bSong\s+\d+", re.IGNORECASE),
        "valid_months": {"january", "march", "may", "july", "september", "november"},
        "section_colors": {
            "TREASURES FROM GOD'S WORD": (98, 101, 104),
            "APPLY YOURSELF TO THE FIELD MINISTRY": (189, 142, 22),
        },
        "default_section_color": (148, 54, 52),
    },
    "yoruba": {
        "base_url": "https://wol.jw.org/",
        "workbook_path": (
            "yo/wol/library/r36/lp-yr/gbogbo-%C3%ACt%E1%BA%B9j%C3%A1de/"
            "%C3%ACw%C3%A9-%C3%ACp%C3%A0d%C3%A9/%C3%ACw%C3%A9-%C3%ACp%C3%A0d%C3%A9"
            "-%C3%ACgb%C3%A9s%C3%AD-ay%C3%A9-%C3%A0ti-i%E1%B9%A3%E1%BA%B9"
            "-%C3%B2j%C3%AD%E1%B9%A3%E1%BA%B9-2026/{month}"
        ),
        "song_pattern": re.compile(r"\bOrin\s+\d+", re.IGNORECASE),
        "valid_months": {
            # Yoruba month slugs as they appear in the URL
            "january",
            "february",
            "march",
            "april",
            "may",
            "june",
            "july",
            "august",
            "september",
            "october",
            "november",
            "december",
        },
        "section_colors": {
            # Yoruba section header strings
            "ÀWỌN ÌṢÚRA INÚ Ọ̀RỌ̀ ỌLỌ́RUN": (98, 101, 104),
            "TẸRA MỌ́ IṢẸ́ ÌWÀÁSÙ": (189, 142, 22),
        },
        "default_section_color": (148, 54, 52),
    },
}

# Yoruba "Living as Christians" equivalent — used to toggle the 4-col layout
LIVING_AS_CHRISTIANS_HEADERS = {
    "english": "LIVING AS CHRISTIANS",
    "yoruba": "MÁA HÙWÀ TÓ YẸ KRISTẸNI",  # update if the live page differs
}

# ---------------------------------------------------------------------------
# WEB SCRAPER
# ---------------------------------------------------------------------------

BASE_URL = "https://wol.jw.org"
NUMBERED_ENTRY: re.Pattern = re.compile(r"^(?P<number>\d+)\.\s(?P<value>.*)$")


async def extract_time_duration(element: str):
    match = re.search(r"\((\d+ min\.)\)", element)
    return match.group(0) if match else None


async def get_time_excerpt(element: Tag, song_pattern: re.Pattern):
    text = element.text.strip()
    sibling = element.find_next_sibling()
    p_tag = sibling.find("p") if sibling else None
    time_duration = await extract_time_duration(p_tag.text.strip()) if p_tag else ""
    m = NUMBERED_ENTRY.match(text)
    if not m:
        return text
    number = int(m["number"])
    value = m["value"]
    return number, value + " " + (time_duration or "")


async def _parse_url(url: str):
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as response:
            html = await response.text()
            return BeautifulSoup(html, "html.parser")


async def _extract_weeks_and_content(soup: BeautifulSoup):
    weeks = soup.find_all(class_=["row", "card"])
    cover_page = weeks[0]
    cover_title: str = cover_page.find(class_=["cardTitleBlock"]).text.strip()
    date_list = []
    links_per_page = []
    for week in weeks[1:]:
        links_per_page.append(week.a["href"])
        date_list.append(week.a.find(class_="cardTitleBlock").text.strip())
    return cover_title, date_list, links_per_page


async def _parse_week_content(url: str, song_pattern: re.Pattern):
    parse_dict = {}
    soup = await _parse_url(BASE_URL + url)
    try:
        date = soup.find("header").h1.text
    except AttributeError:
        return None
    try:
        meeting_content = [
            await get_time_excerpt(element, song_pattern)
            for element in soup.find(class_="bodyTxt").find_all(
                class_="du-fontSize--base"
            )
            if not song_pattern.search(element.text.strip())
        ]
    except AttributeError:
        return None
    parse_dict["date"] = date
    parse_dict["main_table"] = meeting_content
    return parse_dict


async def parse_workbook_url(month: str, language: str = "english"):
    config = LANGUAGE_CONFIG[language]
    url = config["base_url"] + config["workbook_path"].format(month=month.lower())
    soup = await _parse_url(url)
    cover_title, dates, links_per_page = await _extract_weeks_and_content(soup)
    song_pattern = config["song_pattern"]
    results = await asyncio.gather(
        *[_parse_week_content(link, song_pattern) for link in links_per_page]
    )
    return cover_title, results


# ---------------------------------------------------------------------------
# ASSIGNMENT PARSER  (unchanged)
# ---------------------------------------------------------------------------

SKIP_WEEK_MARKER = "__SKIP_WEEK__"


def normalize_date(raw: str):
    cleaned = re.sub(r"[^\d/\-–—. ]", " ", raw)
    cleaned = re.sub(r"[\s/\-–—.]+", " ", cleaned).strip()
    parts = cleaned.split()
    if len(parts) >= 3:
        try:
            d, m, y = int(parts[0]), int(parts[1]), int(parts[2])
            if 1 <= d <= 31 and 1 <= m <= 12 and 2000 <= y <= 2100:
                return (d, m, y)
        except ValueError:
            pass
    return None


def parse_assignments(text: str) -> dict:
    result = {}
    current_date = None
    assignment_line = re.compile(r"^\s*\((\d+)\)\s*(.+)$")
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.search(r"WEEK OF", line, re.IGNORECASE):
            if current_date:
                result[current_date] = SKIP_WEEK_MARKER
            continue
        key = normalize_date(line)
        if key:
            current_date = key
            if current_date not in result:
                result[current_date] = {}
            continue
        m = assignment_line.match(line)
        if m and current_date and result.get(current_date) != SKIP_WEEK_MARKER:
            num = int(m.group(1))
            raw_value = m.group(2).strip()
            value = re.sub(
                r"^([A-Z]\.)+[A-Z]?\.{2,}\s*|^[A-Za-z]+\.{2,}\s*", "", raw_value
            ).strip()
            result[current_date][num] = value
    return result


def find_assignments_for_date(scraped_date: str, assignments: dict):
    nums = list(map(int, re.findall(r"\d+", scraped_date)))
    if not nums:
        return None
    for key, value in assignments.items():
        day, month, year = key
        if day in nums:
            return value
    return None


# ---------------------------------------------------------------------------
# DOCX BUILDER
# ---------------------------------------------------------------------------


def rgb_to_hex(rgb_color: tuple):
    return "{:02X}{:02X}{:02X}".format(*rgb_color)


def set_cell_background(cell, rgb_color):
    hex_color = rgb_to_hex(rgb_color)
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement("w:shd")
    cell_shading.set(qn("w:val"), "clear")
    cell_shading.set(qn("w:color"), "auto")
    cell_shading.set(qn("w:fill"), hex_color)
    cell_properties.append(cell_shading)


def add_cell_text(
    cell: _Cell,
    text: str,
    text_color: RGBColor = RGBColor(255, 255, 255),
    font_name: str = "Calibri",
    font_size: int = 12,
    bold: bool = False,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.color.rgb = text_color


def add_custom_row(
    word_table: Table,
    no_of_cols: int,
    divisions: list[float],
) -> Sequence[_Cell]:
    if sum(divisions) > 1:
        raise ValueError(f"Divisions sum > 1: {sum(divisions)}")
    if any(div < 0 for div in divisions):
        raise ValueError("Negative division")
    partitions_idx = []
    cell_begin = 0
    table_row_cells = word_table.add_row().cells
    for count, div in enumerate(divisions):
        cell_end = (
            round(cell_begin + (div * no_of_cols) - 1)
            if count != len(divisions) - 1
            else (no_of_cols - 1)
        )
        table_row_cells[cell_begin].merge(table_row_cells[cell_end])
        partitions_idx.append(cell_begin)
        cell_begin = cell_end + 1
    return [table_row_cells[p] for p in partitions_idx]


def build_document(
    cover_title: str,
    workbook_contents: list,
    assignments: dict,
    language: str = "english",
) -> io.BytesIO:
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    config = LANGUAGE_CONFIG[language]
    section_colors: dict = config["section_colors"]
    default_section_color: tuple = config["default_section_color"]
    living_as_christians_header = LIVING_AS_CHRISTIANS_HEADERS[language]

    partner_assignment_pattern = re.compile(
        "Starting a Con|Following Up|Explaining Your|Making Disciples"
    )
    single_assignment_pattern = re.compile(
        r"Bible Reading|\w+—Imitate \w+|\w+—What \w+ Did|Talk"
    )

    cols = 24

    heading = document.add_heading(
        "AKEJA CONGREGATION MIDWEEK MEETING SCHEDULE", level=1
    )
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(4)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    subheading = document.add_heading(cover_title, level=2)
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading.paragraph_format.space_before = Pt(2)
    subheading.paragraph_format.space_after = Pt(10)

    count = 0
    for record in workbook_contents:
        if record is None:
            continue

        week_assignments = find_assignments_for_date(record["date"], assignments)
        is_skip_week = week_assignments == SKIP_WEEK_MARKER
        name_map: dict = (
            {} if (week_assignments is None or is_skip_week) else week_assignments
        )

        new_table = document.add_table(cols=cols, rows=0)

        hdr_partitions = add_custom_row(new_table, cols, [1])
        date_label = record["date"]
        if is_skip_week:
            date_label += " — WEEK OF REGIONAL CONVENTION"
        add_cell_text(
            cell=hdr_partitions[0],
            text=str(date_label),
            bold=True,
            text_color=RGBColor(0, 0, 0),
        )

        first_row = add_custom_row(new_table, cols, divisions=[0.8, 0.2])
        add_cell_text(
            cell=first_row[0],
            text="Chairman's Opening Comments - 1 min & Prayer",
            text_color=RGBColor(0, 0, 0),
            bold=True,
        )
        add_cell_text(
            cell=first_row[1],
            text="Stephen Adeloro",
            font_size=11,
            text_color=RGBColor(0, 0, 0),
        )

        is_after_living_as_christians = False

        for row in record["main_table"]:
            if row == living_as_christians_header:
                is_after_living_as_christians = True

            if type(row) is not tuple:
                rgb_color = section_colors.get(str(row), default_section_color)
                subheading_row = add_custom_row(new_table, cols, [0.7, 0.3])
                set_cell_background(subheading_row[0], rgb_color=rgb_color)
                add_cell_text(subheading_row[0], str(row), bold=True, font_size=11)
            else:
                row_number = int(row[0])
                assigned_name = name_map.get(row_number, "")
                use_four_cols = not (
                    is_after_living_as_christians or str(row[0]) in {"1", "2"}
                )

                if use_four_cols:
                    meeting_part_row = add_custom_row(
                        new_table, no_of_cols=cols, divisions=[0.1, 0.45, 0.15, 0.3]
                    )
                else:
                    meeting_part_row = add_custom_row(
                        new_table, no_of_cols=cols, divisions=[0.1, 0.6, 0.3]
                    )

                add_cell_text(
                    meeting_part_row[0], str(row[0]), RGBColor(0, 0, 0), font_size=9
                )
                add_cell_text(
                    meeting_part_row[1],
                    str(row[1]),
                    RGBColor(0, 0, 0),
                    font_size=(11 if len(str(row[1])) < 50 else 10),
                )

                if use_four_cols:
                    student_partner_text = (
                        "Student/Partner"
                        if partner_assignment_pattern.match(str(row[1]))
                        else (
                            "Student"
                            if single_assignment_pattern.match(str(row[1]))
                            else ""
                        )
                    )
                    add_cell_text(
                        meeting_part_row[2],
                        text=student_partner_text,
                        text_color=RGBColor(0, 0, 0),
                        font_size=9,
                    )
                    add_cell_text(
                        meeting_part_row[3],
                        text=assigned_name,
                        text_color=RGBColor(0, 0, 0),
                        font_size=10,
                    )
                else:
                    add_cell_text(
                        meeting_part_row[2],
                        text=assigned_name,
                        text_color=RGBColor(0, 0, 0),
                        font_size=10,
                    )

        last_row = add_custom_row(new_table, cols, [0.1, 0.6, 0.3])
        add_cell_text(
            last_row[1],
            text="• Review/Preview/Announcements - (3 min.) ",
            bold=True,
            text_color=RGBColor(0, 0, 0),
            font_size=11,
        )

        count += 1
        if count % 2 == 1:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
        if count % 2 == 0:
            document.add_page_break()

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# ROUTES
# ---------------------------------------------------------------------------

VALID_MONTHS_ENGLISH = {"january", "march", "may", "july", "september", "november"}
# The Yoruba workbook URL uses the same English month slugs
VALID_MONTHS_YORUBA = {"january", "march", "may", "july", "september", "november"}


@app.get("/", response_class=HTMLResponse)
async def index():
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()


@app.post("/generate")
async def generate(
    month: str = Form(...),
    assignments: str = Form(""),
    assignments_file: UploadFile = File(None),
    language: str = Form("english"),
):
    language = language.lower()
    if language not in LANGUAGE_CONFIG:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported language: {language}. Choose 'english' or 'yoruba'.",
        )

    valid_months = LANGUAGE_CONFIG[language]["valid_months"]
    if month.lower() not in valid_months:
        raise HTTPException(
            status_code=400, detail=f"Invalid month for {language}: {month}"
        )

    assignment_text = assignments
    if assignments_file and assignments_file.filename:
        content = await assignments_file.read()
        assignment_text = content.decode("utf-8", errors="ignore")

    parsed_assignments = (
        parse_assignments(assignment_text) if assignment_text.strip() else {}
    )

    try:
        cover_title, workbook_contents = await parse_workbook_url(month, language)
    except Exception as e:
        ic(f"Failed to fetch workbook: {str(e)}")
        raise HTTPException(
            status_code=502,
            detail="Failed to fetch workbook. It seems like there is a problem with your network connection.",
        )

    buffer = build_document(
        cover_title, workbook_contents, parsed_assignments, language
    )

    safe_title = cover_title.encode("latin-1", errors="replace").decode("latin-1")
    filename = f"{safe_title}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/health")
async def health():
    return {"status": "ok"}
