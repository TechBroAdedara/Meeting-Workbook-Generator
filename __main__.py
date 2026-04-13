import asyncio
import re
import time
from collections.abc import Sequence
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Inches
from docx.table import Table, _Cell

from web_scrapper import parse_workbook_url

document = Document()

sections = document.sections
for section in sections:
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


# ---------------------------------------------------------------------------
# ASSIGNMENT PARSER
# ---------------------------------------------------------------------------

SKIP_WEEK_MARKER = "__SKIP_WEEK__"


def normalize_date(raw: str) -> tuple[int, int, int] | None:
    """
    Accepts messy date strings like:
        '04/05/2026', '18-----05-------2026', 'OCTOBER 27–NOVEMBER 2'
    and returns a (day, month, year) int tuple, or None if unparseable.
    """
    # Strip everything that isn't a digit or common separator
    cleaned = re.sub(r"[^\d/\-–—. ]", " ", raw)
    # Collapse runs of separators / spaces into a single space
    cleaned = re.sub(r"[\s/\-–—.]+", " ", cleaned).strip()

    parts = cleaned.split()
    if len(parts) >= 3:
        try:
            # Handles DD MM YYYY (most common in your data)
            d, m, y = int(parts[0]), int(parts[1]), int(parts[2])
            if 1 <= d <= 31 and 1 <= m <= 12 and 2000 <= y <= 2100:
                return (d, m, y)
        except ValueError:
            pass

    return None


def parse_assignments(text: str) -> dict[tuple[int, int, int], dict[int, str] | str]:
    """
    Parses your hand-typed assignment block into a dict:

        {
            (4, 5, 2026): {3: "B.R...Bassey Ubong", 4: "S.C...Tracy/Elizabeth", ...},
            (18, 5, 2026): "__SKIP_WEEK__",
            ...
        }

    Lines that look like 'XX/XX/XXXX' or similar start a new date block.
    Lines that look like '(N) ...' are assignment entries for that block.
    Lines containing 'WEEK OF' trigger a skip marker.
    Month/year-only header lines (e.g. 'MAY 2026') are ignored.
    """
    result: dict[tuple[int, int, int], dict[int, str] | str] = {}
    current_date: tuple[int, int, int] | None = None

    # Pattern: optional leading whitespace, (number) rest-of-line
    assignment_line = re.compile(r"^\s*\((\d+)\)\s*(.+)$")

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # --- Skip-week marker ---
        if re.search(r"WEEK OF", line, re.IGNORECASE):
            if current_date:
                result[current_date] = SKIP_WEEK_MARKER
            continue

        # --- Try to parse as a date header ---
        key = normalize_date(line)
        if key:
            current_date = key
            if current_date not in result:
                result[current_date] = {}
            continue

        # --- Try to parse as an assignment line ---
        m = assignment_line.match(line)
        if m and current_date and result.get(current_date) != SKIP_WEEK_MARKER:
            num = int(m.group(1))
            raw_value = m.group(2).strip()
            # Strip prefix like "B.R...", "S.C....", "E.Y.B...", "Talk...." etc.
            # A prefix is: single letters separated by dots (B.R / E.Y.B)
            # OR a plain word (Talk, Follow), followed by two or more dots.
            # The mandatory "..." separator is what distinguishes it from a real name.
            value = re.sub(r"^([A-Z]\.)+[A-Z]?\.{2,}\s*|^[A-Za-z]+\.{2,}\s*", "", raw_value).strip()
            result[current_date][num] = value  # type: ignore[index]

    return result


def find_assignments_for_date(
    scraped_date: str,
    assignments: dict[tuple[int, int, int], dict[int, str] | str],
) -> dict[int, str] | str | None:
    """
    Tries to match a scraped week date string (e.g. 'May 4–10')
    against the keys in the assignments dict.

    Strategy: extract any numbers from the scraped date, try each candidate
    in the assignments dict whose day matches any of those numbers and whose
    month matches the scraped month if detectable.
    """
    # Extract all integers from the scraped date
    nums = list(map(int, re.findall(r"\d+", scraped_date)))
    if not nums:
        return None

    # Try direct day match first
    for key, value in assignments.items():
        day, month, year = key
        if day in nums:
            return value

    return None


# ---------------------------------------------------------------------------
# HELPERS (unchanged from original)
# ---------------------------------------------------------------------------

def rgb_to_hex(rgb_color: tuple):
    return "{:02X}{:02X}{:02X}".format(*rgb_color)


def set_cell_background(cell, rgb_color):
    hex_color = rgb_to_hex(rgb_color)
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement("w:shd")
    cell_shading.set(qn("w:val"), "clear")
    cell_shading.set(qn("w:color"), "auto")
    cell_shading.set(qn("w:fill"), hex_color)
    cell_properties.append(cell_shading)


def add_cell_text(
    cell: _Cell,
    text: str,
    text_color: RGBColor = RGBColor(255, 255, 255),
    font_name: str = "Calibri",
    font_size: int = 12,
    bold: bool = False,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.color.rgb = text_color


def add_custom_row(
    word_table: Table,
    no_of_cols: int,
    divisions: list[float],
) -> Sequence[_Cell]:
    if sum(divisions) > 1:
        raise ValueError(
            f"The sum of divisions should not be greater than 1 (100%). Current sum: {sum(divisions)} "
        )
    if any([div < 0 for div in divisions]):
        raise ValueError("You cannot have a negative division")

    partitions_idx = []
    cell_begin = 0
    table_row_cells = word_table.add_row().cells

    count = 0
    for div in divisions:
        cell_end = (
            round(cell_begin + (div * no_of_cols) - 1)
            if count != len(divisions) - 1
            else (no_of_cols - 1)
        )
        table_row_cells[cell_begin].merge(table_row_cells[cell_end])
        partitions_idx.append(cell_begin)
        cell_begin = cell_end + 1
        count += 1

    return [table_row_cells[partition] for partition in partitions_idx]


# ---------------------------------------------------------------------------
# MAIN GENERATOR
# ---------------------------------------------------------------------------

async def generate_word_document():
    cols = 24

    partner_assignment_pattern = re.compile(
        "Starting a Con|Following Up|Explaining Your|Making Disciples"
    )
    single_assignment_pattern = re.compile(
        r"Bible Reading|\w+—Imitate \w+|\w+—What \w+ Did|Talk"
    )

    start = time.perf_counter()

    # --- Month input ---
    month: Literal["January", "March", "May", "July", "September", "November"] = input(  # type: ignore
        "Enter the month you want to get workbook contents for?"
        ' \n Valid months: ["January", "March", "May", "July", "September", "November"] \n'
    )
    list_of_possible_months = ["january", "march", "may", "july", "september", "november"]
    if month.lower() not in list_of_possible_months:
        raise ValueError("Invalid month. Please input a month from the valid list of months.")

    # --- Assignment file input ---
    assignment_file = input(
        "\nEnter the path to your assignments .txt file "
        "(or press Enter to skip and leave names blank): "
    ).strip()

    assignments = {}
    if assignment_file:
        try:
            with open(assignment_file, "r", encoding="utf-8") as f:
                assignment_text = f.read()
            assignments = parse_assignments(assignment_text)
            print(f"Parsed assignments for {len(assignments)} week(s).")
        except FileNotFoundError:
            print(f"File not found: {assignment_file!r} — continuing without assignments.")
    else:
        print("No assignments file provided — last column will be left blank.")

    # --- Web scrape ---
    cover_title, workbook_contents = await parse_workbook_url(month)

    total_progress = len(workbook_contents)
    count = 0

    print("\n\n")
    print("<-----------------WORD DOCUMENT PROCESSING----------------------->")

    heading = document.add_heading("AKEJA CONGREGATION MIDWEEK MEETING SCHEDULE", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(4)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    subheading = document.add_heading(cover_title, level=2)
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading.paragraph_format.space_before = Pt(2)
    subheading.paragraph_format.space_after = Pt(10)

    for record in workbook_contents:
        if record is None:
            print(f"SKIPPED : {record}")
            continue

        # Look up assignments for this week
        week_assignments = find_assignments_for_date(record["date"], assignments)
        is_skip_week = week_assignments == SKIP_WEEK_MARKER
        name_map: dict[int, str] = {} if (week_assignments is None or is_skip_week) else week_assignments  # type: ignore

        new_table = document.add_table(cols=cols, rows=0)

        # Header row (date)
        hdr_partitions = add_custom_row(new_table, cols, [1])
        date_label = record["date"]
        if is_skip_week:
            date_label += " — WEEK OF REGIONAL CONVENTION"
        add_cell_text(
            cell=hdr_partitions[0],
            text=str(date_label),
            bold=True,
            text_color=RGBColor(0, 0, 0),
        )

        # First static row
        first_row = add_custom_row(new_table, cols, divisions=[0.8, 0.2])
        add_cell_text(
            cell=first_row[0],
            text="Chairman's Opening Comments - 1 min & Prayer",
            text_color=RGBColor(0, 0, 0),
            bold=True,
        )
        add_cell_text(
            cell=first_row[1],
            text="Stephen Adeloro",
            font_size=11,
            text_color=RGBColor(0, 0, 0),
        )

        is_after_living_as_christians = False

        for row in record["main_table"]:
            if row == "LIVING AS CHRISTIANS":
                is_after_living_as_christians = True

            if type(row) is not tuple:
                # Section header row (Treasures / Apply Yourself / Living as Christians)
                rgb_color = (
                    (98, 101, 104)
                    if str(row) == "TREASURES FROM GOD'S WORD"
                    else (
                        (189, 142, 22)
                        if str(row) == "APPLY YOURSELF TO THE FIELD MINISTRY"
                        else (148, 54, 52)
                    )
                )
                subheading_row = add_custom_row(new_table, cols, [0.7, 0.3])
                set_cell_background(subheading_row[0], rgb_color=rgb_color)
                add_cell_text(subheading_row[0], str(row), bold=True, font_size=11)

            else:
                # Assignment row
                row_number = int(row[0])
                assigned_name = name_map.get(row_number, "")

                # Decide column layout:
                # 4 cols if Apply Yourself part (not living-as-christians, not rows 1/2)
                # 3 cols otherwise (no student/partner label column)
                use_four_cols = not (
                    is_after_living_as_christians or str(row[0]) in {"1", "2"}
                )

                if use_four_cols:
                    # [num | title | student label | name]
                    meeting_part_row = add_custom_row(
                        new_table, no_of_cols=cols, divisions=[0.1, 0.45, 0.15, 0.3]
                    )
                else:
                    # [num | title | name]
                    meeting_part_row = add_custom_row(
                        new_table, no_of_cols=cols, divisions=[0.1, 0.6, 0.3]
                    )

                # Col 0 – row number
                add_cell_text(
                    meeting_part_row[0], str(row[0]), RGBColor(0, 0, 0), font_size=9
                )

                # Col 1 – assignment title
                add_cell_text(
                    meeting_part_row[1],
                    str(row[1]),
                    RGBColor(0, 0, 0),
                    font_size=(11 if len(str(row[1])) < 50 else 10),
                )

                if use_four_cols:
                    # Col 2 – student/partner label
                    student_partner_text = (
                        "Student/Partner"
                        if partner_assignment_pattern.match(str(row[1]))
                        else (
                            "Student"
                            if single_assignment_pattern.match(str(row[1]))
                            else ""
                        )
                    )
                    add_cell_text(
                        meeting_part_row[2],
                        text=student_partner_text,
                        text_color=RGBColor(0, 0, 0),
                        font_size=9,
                    )
                    # Col 3 – assigned name
                    add_cell_text(
                        meeting_part_row[3],
                        text=assigned_name,
                        text_color=RGBColor(0, 0, 0),
                        font_size=10,
                    )
                else:
                    # Col 2 – assigned name
                    add_cell_text(
                        meeting_part_row[2],
                        text=assigned_name,
                        text_color=RGBColor(0, 0, 0),
                        font_size=10,
                    )

        # Last static row
        last_row = add_custom_row(new_table, cols, [0.1, 0.6, 0.3])
        add_cell_text(
            last_row[1],
            text="• Review/Preview/Announcements - (3 min.) ",
            bold=True,
            text_color=RGBColor(0, 0, 0),
            font_size=11,
        )

        count += 1

        if count % 2 == 1:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)

        if count % 2 == 0:
            document.add_page_break()

        print(
            f"\rProcessed word document for {record['date']} - Progress: {round((count / total_progress) * 100)}%",
            end=" ",
            flush=True,
        )

    print("\n\nWord document has finished processing.")
    end = time.perf_counter()
    print(f"Ran for {end - start:.4f} seconds")

    document.save(f"GeneratedDocs/{cover_title}.docx")


if __name__ == "__main__":
    asyncio.run(generate_word_document())